import csv
import io
import json
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional


@dataclass
class ProcessedFile:
    """Result of processing an uploaded file."""

    filename: str
    file_type: str
    content_type: str
    text_content: Optional[str] = None
    structured_data: Optional[List[Dict[str, Any]]] = None
    row_count: int = 0
    error: Optional[str] = None


class FileProcessor:
    """Process various file formats to extract content."""

    SUPPORTED_EXTENSIONS = {
        ".txt": "text",
        ".md": "text",
        ".log": "text",
        ".csv": "csv",
        ".json": "json",
        ".xml": "xml",
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
    }

    def process_file(self, filename: str, content: bytes) -> ProcessedFile:
        """Process a file and extract its content."""
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return ProcessedFile(
                filename=filename,
                file_type=ext,
                content_type="unknown",
                error=(
                    f"Unsupported file type: {ext}. Supported: "
                    f"{', '.join(sorted(self.SUPPORTED_EXTENSIONS.keys()))}"
                ),
            )

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        try:
            if file_type == "text":
                return self._process_text(filename, ext, content)
            if file_type == "csv":
                return self._process_csv(filename, content)
            if file_type == "json":
                return self._process_json(filename, content)
            if file_type == "xml":
                return self._process_xml(filename, content)
            if file_type == "pdf":
                return self._process_pdf(filename, content)
            if file_type == "docx":
                return self._process_docx(filename, content)
            if file_type == "xlsx":
                return self._process_xlsx(filename, content)
            return ProcessedFile(
                filename=filename,
                file_type=ext,
                content_type="unknown",
                error=f"Processor not implemented for: {file_type}",
            )
        except Exception as e:
            return ProcessedFile(
                filename=filename,
                file_type=ext,
                content_type="error",
                error=f"Processing error: {str(e)}",
            )

    def _decode_text(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")

    def _process_text(self, filename: str, ext: str, content: bytes) -> ProcessedFile:
        """Process plain text files."""
        text = self._decode_text(content)
        return ProcessedFile(
            filename=filename,
            file_type=ext,
            content_type="text",
            text_content=text,
            row_count=len(text.splitlines()),
        )

    def _process_csv(self, filename: str, content: bytes) -> ProcessedFile:
        """Process CSV files."""
        text = self._decode_text(content)
        reader = csv.DictReader(io.StringIO(text))
        rows: List[Dict[str, Any]] = list(reader)

        text_content = self._structured_to_text(rows)

        return ProcessedFile(
            filename=filename,
            file_type=".csv",
            content_type="structured",
            text_content=text_content,
            structured_data=rows,
            row_count=len(rows),
        )

    def _process_json(self, filename: str, content: bytes) -> ProcessedFile:
        """Process JSON files."""
        text = self._decode_text(content)
        data = json.loads(text)

        if isinstance(data, list):
            rows = [r for r in data if isinstance(r, dict)]
            if not rows and data:
                text_content = text
            else:
                text_content = (
                    self._structured_to_text(rows) if rows else text
                )
            return ProcessedFile(
                filename=filename,
                file_type=".json",
                content_type="structured",
                text_content=text_content,
                structured_data=rows or None,
                row_count=len(rows),
            )

        if isinstance(data, dict):
            rows: List[Dict[str, Any]] = []
            for key in ("bugs", "issues", "items", "data", "records", "results"):
                if key in data and isinstance(data[key], list):
                    rows = [r for r in data[key] if isinstance(r, dict)]
                    break
            else:
                rows = [data]

            text_content = self._structured_to_text(rows) if rows else text
            return ProcessedFile(
                filename=filename,
                file_type=".json",
                content_type="structured",
                text_content=text_content,
                structured_data=rows,
                row_count=len(rows),
            )

        rows: List[Dict[str, Any]] = []
        return ProcessedFile(
            filename=filename,
            file_type=".json",
            content_type="structured",
            text_content=text,
            structured_data=rows or None,
            row_count=0,
        )

    def _process_xml(self, filename: str, content: bytes) -> ProcessedFile:
        """Process XML files (including Jira exports)."""
        text = self._decode_text(content)
        root = ET.fromstring(text)
        rows = self._xml_to_dicts(root)

        text_content = self._structured_to_text(rows) if rows else text

        return ProcessedFile(
            filename=filename,
            file_type=".xml",
            content_type="structured",
            text_content=text_content,
            structured_data=rows,
            row_count=len(rows),
        )

    def _xml_to_dicts(self, root: ET.Element) -> List[Dict[str, Any]]:
        """Convert XML tree to list of dictionaries."""
        results: List[Dict[str, Any]] = []

        issue_tags = ["item", "issue", "bug", "defect", "entry", "record", "row"]

        for tag in issue_tags:
            items = root.findall(f".//{tag}")
            if items:
                for item in items:
                    row = self._element_to_dict(item)
                    if row:
                        results.append(row)
                break

        if not results:
            single = self._element_to_dict(root)
            if single:
                results = [single]

        return results

    def _element_to_dict(self, element: ET.Element) -> Dict[str, Any]:
        """Convert XML element to dictionary."""
        result: Dict[str, Any] = {}
        result.update(element.attrib)

        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text or ""
            else:
                result[child.tag] = self._element_to_dict(child)

        if element.text and element.text.strip():
            result["_text"] = element.text.strip()

        return result

    def _structured_to_text(
        self, rows: List[Dict[str, Any]], max_rows: int = 100
    ) -> str:
        """Convert structured data to readable text for AI analysis."""
        if not rows:
            return ""

        lines: List[str] = []
        for i, row in enumerate(rows[:max_rows]):
            lines.append(f"--- Record {i + 1} ---")
            if isinstance(row, dict):
                for key, value in row.items():
                    if value and str(value).strip():
                        lines.append(f"{key}: {value}")
            else:
                lines.append(str(row))
            lines.append("")

        if len(rows) > max_rows:
            lines.append(f"... and {len(rows) - max_rows} more records")

        return "\n".join(lines)

    def _process_pdf(self, filename: str, content: bytes) -> ProcessedFile:
        """Process PDF files."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts: List[str] = []

            for page in reader.pages:
                pg_text = page.extract_text()
                if pg_text:
                    text_parts.append(pg_text)

            full_text = "\n\n".join(text_parts)

            return ProcessedFile(
                filename=filename,
                file_type=".pdf",
                content_type="text",
                text_content=full_text,
                row_count=len(reader.pages),
            )
        except ImportError:
            return ProcessedFile(
                filename=filename,
                file_type=".pdf",
                content_type="error",
                error="PDF support requires pypdf: pip install pypdf",
            )

    def _process_docx(self, filename: str, content: bytes) -> ProcessedFile:
        """Process Word documents."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts: List[str] = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)

            return ProcessedFile(
                filename=filename,
                file_type=".docx",
                content_type="text",
                text_content=full_text,
                row_count=len(text_parts),
            )
        except ImportError:
            return ProcessedFile(
                filename=filename,
                file_type=".docx",
                content_type="error",
                error="Word support requires python-docx: pip install python-docx",
            )

    def _process_xlsx(self, filename: str, content: bytes) -> ProcessedFile:
        """Process Excel files."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            all_rows: List[Dict[str, Any]] = []

            for sheet in wb.worksheets:
                sheet_rows = list(sheet.iter_rows(values_only=True))
                if not sheet_rows:
                    continue

                headers = [
                    str(h) if h is not None else f"col_{i}"
                    for i, h in enumerate(sheet_rows[0])
                ]

                for row in sheet_rows[1:]:
                    if any(cell is not None for cell in row):
                        row_dict = {
                            headers[i]: str(cell) if cell is not None else ""
                            for i, cell in enumerate(row)
                            if i < len(headers)
                        }
                        all_rows.append(row_dict)

            text_content = self._structured_to_text(all_rows)

            return ProcessedFile(
                filename=filename,
                file_type=".xlsx",
                content_type="structured",
                text_content=text_content,
                structured_data=all_rows,
                row_count=len(all_rows),
            )
        except ImportError:
            return ProcessedFile(
                filename=filename,
                file_type=".xlsx",
                content_type="error",
                error="Excel support requires openpyxl: pip install openpyxl",
            )


file_processor = FileProcessor()
